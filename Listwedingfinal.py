import streamlit as st
import pandas as pd
import os
from datetime import datetime
from io import BytesIO

# Docx imports with error handling
try:
    import docx
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

st.set_page_config(page_title="💍 Wedding Guest Manager", layout="wide")
DATA_FILE = "guest_list.csv"

# ---------------- FULL NAMES LIST ----------------
all_names = [
    "Adv. Altaf Junejo Sindh Bar Council",
    "Adv. Bhagwan Das Bheel",
    "Adv. Bhagwandas Bheel",
    "Adv. Pithumal Parwano",
    "Adv. Velji Rathore",
    "Adv. Abdul Jabbar Abbassi",
    "Adv. Abdul Qayoom Malkani",
    "Adv. Abdul Razzak Dars",
    "Adv. Abdul Sattar Bajeer",
    "Adv. Abdullah Laghari",
    "Adv. Abid Chang",
    "Adv. Adam Rajar",
    "Adv. Aftab Ghori",
    "Adv. Ameer Hamzo",
    "Adv. Anad Gonh",
    "Adv. Anchalaram",
    "Adv. Aneel Kumar Rathore",
    "Adv. Aneel Kumar Rathore a",
    "Adv. Arif Kellar",
    "Adv. Asad Ali Murree",
    "Adv. Asad Memon",
    "Adv. Ashaque Ali Bajeer",
    "Adv. Ashok Kumar hyd",
    "Adv. Atta Chandio",
    "Adv. Bhagchand Bheel",
    "Adv. Bharat Heemrani",
    "Adv. Bhawan Kumar MPK",
    "Adv. Bhopat Kolhi",
    "Adv. Bhoro Kolhi",
    "Adv. Bhrulal Heemarani",
    "Adv. Chaman Lal",
    "Adv. Chander Kumar Kolhi",
    "Adv. Chander Tharani",
    "Adv. Dasrat kumar Sukhani",
    "Adv. Dayal Das",
    "Adv. Dayal Das Meghwar",
    "Adv. Devat Rai",
    "Adv. Dhanraj Palani",
    "Adv. Dileep Kumar Karachi",
    "Adv. Dileep kummar Karachi",
    "Adv. Faqeer Mohd Laghari",
    "Adv. Faqeer Munwar Sagar",
    "Adv. Faresh Kumar",
    "Adv. Ghamoon Mal",
    "Adv. Ghulam Hussain Palari",
    "Adv. Ghulamullah Abro",
    "Adv. Govind Meghwar",
    "Adv. Goving Mehraj",
    "Adv. Gulab Meghwar",
    "Adv. Hameed laghari",
    "Adv. Hamir ji Bheel",
    "Adv. Hamlesh Suthar",
    "Adv. Haroon Keerio",
    "Adv. Hashim laghari",
    "Adv. Hassan Jalalani",
    "Adv. Heemraj Bheel",
    "Adv. Hemandas Sanghani",
    "Adv. Hotchand Togani",
    "Adv. Imam Bux Dars",
    "Adv. Imran Memon Adpp",
    "Adv. Inderjeet Lohano",
    "Adv. Islamuddin Rahmoon",
    "Adv. Jai Dev Suthar",
    "Adv. Jai Ki Rai",
    "Adv. Jairam Das",
    "Adv. Jameel Khanzado",
    "Adv. Javaid Akhtar",
    "Adv. Jethanand",
    "Adv. Kabeer Rajar",
    "Adv. Kaleemullah",
    "Adv. Kanji Mal",
    "Adv. Kanji Mal ASC MPK",
    "Adv. Kanwar Amar",
    "Adv. Kelash Kolhi",
    "Adv. Kewal Bheel",
    "Adv. Kewal Gomani",
    "Adv. Khaleel Laghari",
    "Adv. Kuldeep Sharma",
    "Adv. Kundhan Malhi",
    "Adv. Lajpat Rai Panjwani",
    "Adv. Lajpat Rai Soorani",
    "Adv. Leela Ram",
    "Adv. Leelaram",
    "Adv. Love Kumar",
    "Adv. Maeen Bajeer",
    "Adv. Mahesh Bheel",
    "Adv. Majid Aziz",
    "Adv. Malook Khaskhali",
    "Adv. Mangal Meghwar",
    "Adv. Mir Moh Samejo",
    "Adv. Mir Sajid Khoso",
    "Adv. Mohammad Khan Bajeer",
    "Adv. Mohan Lal Bheel",
    "Adv. Mohan Lal Manthrani",
    "Adv. Mohan Lal Rathore MPK",
    "Adv. Moula Bux Bajeer",
    "Adv. Muhammad Amir Qureshi",
    "Adv. Muhammad Bajeer",
    "Adv. Muhammad Hingorjo",
    "Adv. Muhammad Ramzan Chandio",
    "Adv. Mukesh Kumar Rathore",
    "Adv. Mumtaz Jarwar",
    "Adv. Muneer Gilal",
    "Adv. Murtaza Keerio",
    "Adv. Mustafa Hingorjo",
    "Adv. Nabi bux Inspector",
    "Adv. Nadeem Tagar",
    "Adv. Naeem Gddai",
    "Adv. Naeem Talpur APG",
    "Adv. Naveed Jarwar s",
    "Adv. Neel Parkash MPK",
    "Adv. Noor Ahmed Soomro",
    "Adv. Om Parkash",
    "Adv. Parkash Kumar",
    "Adv. Partab Rai",
    "Adv. Permanand",
    "Adv. Pirkash Singh",
    "Adv. Preet Pal Singh",
    "Adv. Qamar Nohri",
    "Adv. Qamat Rai",
    "Adv. Qurban Ali Samejo",
    "Adv. Rafique Hingorjo",
    "Adv. Rai Singh Sodho",
    "Adv. Rajesh Kumar",
    "Adv. Rajesh Kumar Uk",
    "Adv. Ramchand Archna",
    "Adv. Ramesh Kumar",
    "Adv. Ramesh Kumar Depal",
    "Adv. Ramesh Kumar Oad",
    "Adv. Ranghan",
    "Adv. Ravi Kumar",
    "Adv. Riaz Ali Panhwar",
    "Adv. Rizwan Memon",
    "Adv. Roopchan Ragastani",
    "Adv. Roopl Mala Singh",
    "Adv. Saffar Khokhar",
    "Adv. Samiullah Abbassi",
    "Adv. Sanaullah Jhetiyal",
    "Adv. Santosh Kolhi",
    "Adv. Santosh Kumar",
    "Adv. Sarang Ram",
    "Adv. Sawan Sindhi, Mr Sikandar , Adv. Khushal Das",
    "Adv. Shabeer Arbab Hala",
    "Adv. Shahbaz Khaskhali",
    "Adv. Shanker Lal Meghwar",
    "Adv. Shanker Lal Rathore",
    "Adv. Shanker Suthar",
    "Adv. Sheroz Chang",
    "Adv. Shewak Rathore",
    "Adv. Shoukat Sindhi",
    "Adv. Sooba Bhatti",
    "Adv. Sooraj Kumar Qambar",
    "Adv. Subash Sharma",
    "Adv. Subhan Samejo",
    "Adv. Suneel Kumar Goswami",
    "Adv. Suneel Parhiyar",
    "Adv. Suneel Shatrogan",
    "Adv. Taimor Ali Shah",
    "Adv. Tano Mal",
    "Adv. Teerath kumar Jhangi",
    "Adv. Vasand Thari",
    "Adv. Vikram Kumar",
    "Adv. Vikram Meshwari",
    "Adv. Vishal",
    "Adv. Vishandas",
    "Adv. Wahab Channa",
    "Adv. Wahid Khan Adpp",
    "Adv. Waqar Dda",
    "Adv. Wazeer Chandio",
    "Adv. Yasir Khoso",
    "Adv. Zaheer Nohri",
    "Adv. Zaheeruddin Junejo",
    "Adv. Zahid Shah",
    "Adv. Zakaullah Bajeer",
    "Dr. Bharat Lal Dow Internal University",
    "Dr. Jeewat Rai Assistant Chief Planning & Development Department",
    "Dr. Khanji Mal",
    "Dr. Khatu Mal Jewan MPA",
    "Dr. Lal Malkani",
    "Dr. Raja",
    "Dr. Sagram Das (Mehmoodabad)",
    "Engr. Damoon Mal",
    "Engr. Gulji",
    "Engr. Khaku Mal Assistant Professor IBA  (Campus)",
    "Engr. Mithu Mal",
    "Engr. Vadhu Ex. Cheif Engineer",
    "Engr. Veenjho Mal",
    "Mr. & Mrs. Rawal Memon CEO Awaz TV",
    "Mr. & Mrs. Sharjeel Inam Memon  (Senior Minister Information & Transport Mass Transit Department Government of Sindh)",
    "Mr. Abdul Qayoom Mehar",
    "Mr. Abdul Rasheed Khoso",
    "Mr. Abdul Rasheed Solangi  (Secretary Government of Sindh)",
    "Mr. Abdul Sattar Section Officer, Finance Department",
    "Mr. Abdul Sattar Shaikh AC",
    "Mr. Abdul shakoor",
    "Mr. Abdullah Channa",
    "Mr. Abdullah Meo",
    "Mr. Abdullah Shoro",
    "Mr. Adeeb Book Depo",
    "Adv. Raichand Harijan MPK",
    "Mr. Aftab Abbassi",
    "Mr. Ahmed Ali Solangi",
    "Mr. Ahmed Kalroo PS K.D Jewan",
    "Mr. Akash Kumar Judge",
    "Mr. Akash Oad",
    "Mr. Akash Santorai",
    "Mr. Akhlaque Hussain Larik",
    "Mr. Ali Hyder AC",
    "Mr. Altaf Dimro State life",
    "Mr. Amar Lal",
    "Mr. Amar Lal Oad",
    "Mr. Amjad Tunio",
    "Mr. Amolakh",
    "Mr. Amolakh Das Uk",
    "Mr. Amrat Lal Punhnani",
    "Mr. Asghar Jalalani",
    "Mr. Ashok khatwani",
    "Mr. Ashok kumar diplo",
    "Mr. Atif Bhatti  (Director Admin Excise Dept)",
    "Mr. Attam Kumar",
    "Mr. Audu Mal",
    "Mr. Babar Book Depo",
    "Mr. Basheer Bhutto PS Planning & Development Department",
    "Mr. Bhadu Mal Mr. Amolakh Das Mr. Chandar Lal, Mr. Ramesh Kumar (Roopchand) and Mr. Bhuro Lal",
    "Mr. Bhagat Arib, Mr. Mendharo Mal and Mr. Om Parkish",
    "Mr. Bhagchand House Building",
    "Mr. Bhemraj Lecturer",
    "Mr. Bheru Mal Balani Ex. MPA",
    "Mr. Bhopat Rai Senior Civil Judge",
    "Mr. Bhuralal Chaglani",
    "Mr. Chahno Mal House Building",
    "Mr. Chander Kanji mal",
    "Mr. Chando Mal",
    "Mr. Chatro Assistant Mukhtiarkar",
    "Mr. Danish Shaikh",
    "Mr. Dasrat Kumar",
    "Mr. Deedar Assistant Finance Department",
    "Mr. Devan Inspector",
    "Mr. Devan Service hospital",
    "Mr. Dewan Teekchand",
    "Mr. Dodo Khan Jatio",
    "Mr. Dolut Ram",
    "Mr. Doongar Doothi",
    "Mr. Dyal Das Deputy Secretary Sindh Secretariat",
    "Mr. Ferhan Memon SCO, Finance Department",
    "Mr. Gautam Rai Education Works (Highway)",
    "Mr. Ghansham Chaglani",
    "Mr. Gordhan NBP",
    "Mr. Govinda Rwf",
    "Mr. Gul Bahar Oad",
    "Mr. Gulab Rai MRI",
    "Mr. Gulji Inspector",
    "Mr. Hafeez",
    "Mr. Hafeez Sindhi",
    "Mr. Hajan Abro",
    "Mr. Haji Ameer Jatt",
    "Mr. Haji Meo",
    "Mr. Hakeem Magrio",
    "Mr. Harchand, Mr. Jai Ram and Mr. Dolat Ram",
    "Mr. Haresh Kumar",
    "Mr. Harichand",
    "Mr. Harji Excise Department",
    "Mr. Hasrat Parkash",
    "Mr. Hassan Ali Durani",
    "Mr. Hoat LUMHS",
    "Mr. Hussain Mansoor  (PRO Senior Minister Sindh)",
    "Mr. Imran Shaikh",
    "Mr. Imtiaz Bhutto",
    "Mr. Imtiaz Khoso",
    "Mr. Inam Kandhro",
    "Mr. Indar Kumar  Ladhani, and Mr. Love Kumar Ladhani",
    "Mr. Irfan Qureshi",
    "Mr. Irjan Das Director Mining & Cool",
    "Mr. Ishaque Channa",
    "Mr. Ishtiaque Tower Market",
    "Mr. Ishtiaque tower mrkt",
    "Mr. Izhar Hussain Buriro  (PS to Senior Minister Sindh)",
    "Mr. Jabbar Shaikh",
    "Mr. Jai Sooraj Civil Judge",
    "Mr. Jai Sooraj Judge",
    "Mr. Jairam Das Kasbo",
    "Mr. Javed ali",
    "Mr. Javeed Ali Phulpoto SCO, Finance Department",
    "Mr. Jeevat Rai SPSC",
    "Mr. Junaid Iqbal Qureshi",
    "Mr. Kami (Shahid Jatt)",
    "Mr. Kamran Magsi",
    "Mr. Kamran Memon  (SMTA Karachi)",
    "Mr. Kanji Mal, Mr. Pervaiz, Mr. Aneel Kumar, Mr. Amrat Lal and Mr. Jai Parkish",
    "Mr. Kashif Samoon",
    "Mr. Kazi Abid Asid (Daily Ibrat)",
    "Mr. Khenraj, Mr. Total Ram and Mr. Karmoon Mal",
    "Mr. Khursheed Meo",
    "Mr. Khushal",
    "Mr. Khushal Sangrish (Manager Sindh Bank)",
    "Mr. Kirshan PPHI",
    "Mr. Kirshan Lal, Mr. Chandar Lal , Mr. Tota Ram, Mr. Dileep Kumar, Mr. Vinod Kumar, Mr. Satram Das , and Mr. Ramesh Kumar",
    "Mr. Lachaman, Mr. Mevo Mal",
    "Mr. Lekhraj",
    "Mr. Long",
    "Mr. Long (Badar Commercial)",
    "Mr. Majid Ali Soomro",
    "Mr. Malik Mulazam Hussain",
    "Mr. Malik Qasir",
    "Mr. Malik Sarfraz",
    "Mr. Malik Shahzad",
    "Mr. Malik Sohail",
    "Mr. Malik Sajjid",
    "Mr. Manoj Kumar",
    "Mr. Mendharo Mal Valasi",
    "Mr. Mendharo Mal, Mr. Kamlesh Kumar and Mr. Rajesh Kumar",
    "Mr. Meva Ram, Mr. Girdhari Lal, Mr. Bharat Lal (Sakri)",
    "Mr. Mevaram Galhoro MPK",
    "Mr. Misri Mal Ladhani Ex Chief Commissioner",
    "Mr. Mohan Lal",
    "Mr. Mohan Lal Jewan (P.E.C.H.S)",
    "Mr. Mohsin Ali",
    "Mr. Moti",
    "Mr. Moula Bux Mirjat",
    "Mr. Muhammad Adeel Soomro Section Officer, Finance Department",
    "Mr. Muhammad Ali Phulphoto SCO, Finance Department",
    "Mr. Muhammad Ali Solangi",
    "Mr. Muhammad Shahid Shaikh",
    "Mr. Muhammad Yameen Abbasi Section Officer, Finance Department",
    "Mr. Mukesh Kumar Khatwani",
    "Mr. Mukesh Kumar Revenue",
    "Mr. Nadeem Buririo Judge",
    "Mr. Nagaram Judge",
    "Mr. Nand Lal Excise Department",
    "Mr. Nand Lal SCO, Finance Department",
    "Mr. Naveed Qureshi",
    "Mr. Nihal Chand South City Hospital",
    "Mr. Nihal Chand SPSC",
    "Mr. Nihal Ex. Director Mr. Amrat Lal",
    "Mr. Nirmal",
    "Mr. Nisar Solangi",
    "Mr. Nizar Imam  Director I.T  Finance Department",
    "Mr. Noor Ahmed Chandio Judge",
    "Mr. Parkash Kumar Judge",
    "Mr. Prem Oad",
    "Mr. Qadar Bux Sahito",
    "Mr. Raheel Abbassi",
    "Mr. Raheel Memon  (Business Man)",
    "Mr. Rahul",
    "Mr. Rahul Khetlarai",
    "Mr. Rajesh Chahoo",
    "Mr. Rajnesh kumar",
    "Mr. Ramesh Kumar",
    "Mr. Ramesh Kumar HESCO",
    "Mr. Rashid Soomro",
    "Mr. Ravi Malkani",
    "Mr. Ravi Oad",
    "Mr. Ravi Shankar FIA",
    "Mr. Ravi Shanker",
    "Mr. Reejhu Mal",
    "Mr. Reejhu Mal Judge",
    "Mr. Revachand",
    "Mr. Sachan PST(Dhabhari)",
    "Mr. Saddar Meo",
    "Mr. Sajjid Nizamani",
    "Mr. Saleem Chandio",
    "Mr. Salman Book Depo",
    "Mr. Salman Mansoor",
    "Mr. Salman Memon",
    "Mr. Sambhu",
    "Mr. Sambhu Arnario",
    "Mr. Santosh Oad",
    "Mr. Sarwan Kumar",
    "Mr. Shahjahan",
    "Mr. Shahzad Qureshi",
    "Mr. Shahzaib Jatt",
    "Mr. Sikandar",
    "Mr. Sonji, Mr. Bhugro and Mr. Raj Kumar",
    "Mr. Sukhram Das",
    "Mr. Surendar Valasai MPA  (Media Cell Bilwal House)",
    "Mr. Suresh Jodhani DDA",
    "Mr. Tanu Mal, Mr. Munesh Kumar and Mr. Attam Kumar",
    "Mr. Taro Mal, Mr. Dolat Ram, Mr. Jahwar, Mr. Sadhat, Mr. Gulshan",
    "Mr. Thakro Mal",
    "Mr. Toto Mal SSGC",
    "Mr. Ubaidullah Nahinoon",
    "Mr. Ugo Mal Deputy District Attorney",
    "Mr. Veenjho",
    "Mr. Vijay Kumar",
    "Mr. Vinod Oad",
    "Mr. Vishan Das",
    "Mr. Vishan Das Manager NBP",
    "Mr. Vishan Das Nahto",
    "Mr. Vishan Das NBP",
    "Mr. Waseem Panhwar",
    "Mr. Wazeer Oad",
    "Mr. Zameer Abbasi  (PD BRT Yellow Line)",
    "Mr. Zeeshan Memon",
    "Mr. Zulfiquar Solangi",
    "Prof. Bheem Raj Kalani",
    "Prof. Hassan Ali Durani",
    "Prof. Mukesh Kumar Khatwani",
    "Prof. Sujo Mal",
    "Seth Nihal Chand",
    "Seth Valam",
    "Seth Valam, Mr Dileep Kumar and Mr. Ashok Kumar",
    "Syed Ghulam Murtaza Finance Department",
    "Syed Abid Shah  (PS to Senior Minister Sindh)",
    "Mr. Zohaib Ahmed Soomro PS Finance Department",
    "Mr. Zahoor Ahmed SCO Finance Department",
    "Mr. Khalil PS Finance Department",
    "Mr. Mashooque Ali President Finance Department",
    "Mr. Wazir Bhand SCO Finance Department"
]

# ---------------- LOAD OR CREATE DATA ----------------
@st.cache_data
def load_data():
    if os.path.exists(DATA_FILE):
        df = pd.read_csv(DATA_FILE)
    else:
        df = pd.DataFrame({
            "SR#": list(range(1, len(all_names)+1)),
            "Name": all_names,
            "City": ["" for _ in all_names],
            "Invited": [False for _ in all_names],
            "Amount": [0 for _ in all_names],
            "Added_Date": [datetime.now().strftime("%Y-%m-%d") for _ in all_names]
        })
        df.to_csv(DATA_FILE, index=False)
    
    # Ensure all names exist and handle NaN values
    existing = df["Name"].tolist()
    for n in all_names:
        if n not in existing:
            df.loc[len(df)] = [len(df)+1, n, "", False, 0, datetime.now().strftime("%Y-%m-%d")]
    
    df["SR#"] = range(1, len(df)+1)
    
    # Fill NaN values
    df['City'] = df['City'].fillna('')
    df['Invited'] = df['Invited'].fillna(False)
    df['Amount'] = df['Amount'].fillna(0)
    
    return df

# Initialize session state
if 'df' not in st.session_state:
    st.session_state.df = load_data()

# Use session state dataframe
df = st.session_state.df

# ---------------- SAVE + REFRESH ----------------
def save_data(updated_df):
    updated_df["SR#"] = range(1, len(updated_df)+1)
    updated_df.to_csv(DATA_FILE, index=False)
    st.session_state.df = updated_df
    st.success("✅ Data Updated!")
    st.rerun()

# ---------------- DOWNLOAD FUNCTIONS ----------------
def create_excel_download(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Guest List')
        
        # Create summary sheet
        summary_data = {
            'Metric': ['Total Guests', 'Invited', 'Not Invited', 'Total Amount (PKR)'],
            'Value': [
                len(df),
                df['Invited'].sum(),
                len(df) - df['Invited'].sum(),
                f"PKR {df['Amount'].sum():,}"
            ]
        }
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_excel(writer, index=False, sheet_name='Summary')
    
    processed_data = output.getvalue()
    return processed_data

def create_word_download(df):
    if not DOCX_AVAILABLE:
        st.error("❌ Word download unavailable. Please install python-docx: pip install python-docx")
        return None
    
    try:
        doc = docx.Document()
        
        # Title
        title = doc.add_heading('Wedding Guest List', 0)
        
        # Summary
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(f'Total Guests: {len(df)}')
        doc.add_paragraph(f'Invited: {df["Invited"].sum()}')
        doc.add_paragraph(f'Not Invited: {len(df) - df["Invited"].sum()}')
        doc.add_paragraph(f'Total Amount: PKR {df["Amount"].sum():,}')
        
        # Guest List
        doc.add_heading('Guest Details', level=1)
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'SR#'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'City'
        hdr_cells[3].text = 'Invited'
        hdr_cells[4].text = 'Amount (PKR)'
        hdr_cells[5].text = 'Added Date'
        
        # Data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['SR#'])
            row_cells[1].text = str(row['Name'])
            row_cells[2].text = str(row['City'])
            row_cells[3].text = 'Yes' if row['Invited'] else 'No'
            row_cells[4].text = f"{row['Amount']:,}"
            row_cells[5].text = str(row.get('Added_Date', ''))
        
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io.getvalue()
    
    except Exception as e:
        st.error(f"Error creating Word document: {e}")
        return None

# ---------------- SIDEBAR ----------------
st.sidebar.title("Navigation")
option = st.sidebar.radio("Select:", [
    "Dashboard", "Add New Guest", "View Guest List", "Search Guest", 
    "Filter & Analytics", "Update Guest", "Remove Guest"
])

# ---------------- DASHBOARD ----------------
if option == "Dashboard":
    st.subheader("📊 Dashboard")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Guests", len(df))
    
    with col2:
        st.metric("Invited", df['Invited'].sum())
    
    with col3:
        st.metric("Not Invited", len(df) - df['Invited'].sum())
    
    with col4:
        total_amount = df['Amount'].sum()
        st.metric("Total Amount", f"PKR {total_amount:,}")

# ---------------- ADD NEW GUEST ----------------
elif option == "Add New Guest":
    st.subheader("➕ Add Guest")
    
    col1, col2 = st.columns(2)
    
    with col1:
        name = st.text_input("Guest Name *")
        city = st.text_input("City/Village")
    
    with col2:
        invited = st.checkbox("Invited ✅")
        amount = st.number_input("Amount (PKR)", min_value=0, step=100, value=0)

    if st.button("Add Guest", type="primary"):
        if name.strip() == "":
            st.warning("⚠ Name cannot be empty!")
        elif name.strip() in df["Name"].values:
            st.error("❌ This name already exists!")
        else:
            new_guest = {
                "SR#": len(df) + 1,
                "Name": name.strip(),
                "City": city.strip(),
                "Invited": invited,
                "Amount": amount,
                "Added_Date": datetime.now().strftime("%Y-%m-%d")
            }
            updated_df = pd.concat([df, pd.DataFrame([new_guest])], ignore_index=True)
            save_data(updated_df)

# ---------------- VIEW LIST ----------------
elif option == "View Guest List":
    st.subheader("📋 Guest List")
    
    # Filters
    col1, col2, col3 = st.columns(3)
    with col1:
        show_invited = st.selectbox("Invitation Status", 
                                  ["All", "Invited Only", "Not Invited"])
    with col2:
        sort_by = st.selectbox("Sort By", 
                              ["SR#", "Name", "City", "Amount"])
    with col3:
        sort_order = st.selectbox("Order", ["Ascending", "Descending"])
    
    # Apply filters
    filtered_df = df.copy()
    if show_invited == "Invited Only":
        filtered_df = filtered_df[filtered_df["Invited"] == True]
    elif show_invited == "Not Invited":
        filtered_df = filtered_df[filtered_df["Invited"] == False]
    
    # Apply sorting
    ascending = sort_order == "Ascending"
    filtered_df = filtered_df.sort_values(by=sort_by, ascending=ascending)
    
    st.write(f"**Total:** {len(filtered_df)} guests")
    st.dataframe(filtered_df, use_container_width=True, hide_index=True)
    
    # Export options
    st.subheader("📥 Export Data")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 Download as Excel"):
            excel_data = create_excel_download(filtered_df)
            st.download_button(
                label="⬇️ Click to Download Excel",
                data=excel_data,
                file_name="wedding_guests.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
    
    with col2:
        if st.button("📝 Download as Word"):
            word_data = create_word_download(filtered_df)
            if word_data is not None:
                st.download_button(
                    label="⬇️ Click to Download Word",
                    data=word_data,
                    file_name="wedding_guests.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("Word download not available")

# ---------------- SEARCH ----------------
elif option == "Search Guest":
    st.subheader("🔍 Search Guest")
    
    search_type = st.radio("Search by:", ["Name", "City", "Both"], horizontal=True)
    query = st.text_input("Enter search term...")
    
    if query:
        # Handle NaN values and case-insensitive search
        if search_type == "Name":
            result = df[df["Name"].str.contains(query, case=False, na=False)]
        elif search_type == "City":
            result = df[df["City"].str.contains(query, case=False, na=False)]
        else:  # Both
            result = df[
                df["Name"].str.contains(query, case=False, na=False) |
                df["City"].str.contains(query, case=False, na=False)
            ]
        
        if len(result) > 0:
            st.write(f"**Found {len(result)} results:**")
            st.dataframe(result, use_container_width=True, hide_index=True)
        else:
            st.info("No guests found matching your search criteria.")
    else:
        st.info("Enter a search term to find guests")

# ---------------- FILTER & ANALYTICS ----------------
elif option == "Filter & Analytics":
    st.subheader("💰 Filter & Analytics")
    
    tab1, tab2, tab3 = st.tabs(["Amount Filter", "City Analysis", "Summary"])
    
    with tab1:
        st.subheader("Filter by Amount")
        
        # Safe amount range calculation
        if df.empty or df['Amount'].isna().all():
            max_amount_val = 10000
            min_amount_val = 0
        else:
            min_amount_val = int(df['Amount'].min())
            max_amount_val = int(df['Amount'].max())
            
            # Handle case where all amounts are the same
            if min_amount_val == max_amount_val:
                max_amount_val = min_amount_val + 10000
        
        # Ensure we have a valid range
        if min_amount_val >= max_amount_val:
            max_amount_val = min_amount_val + 10000
        
        min_amount, max_amount = st.slider(
            "Select Amount Range (PKR)",
            min_value=0,
            max_value=max(10000, max_amount_val),  # Ensure at least 10000 range
            value=(min_amount_val, max_amount_val),
            step=500
        )
        
        filtered = df[(df["Amount"] >= min_amount) & (df["Amount"] <= max_amount)]
        st.write(f"**Guests with amount between {min_amount} and {max_amount} PKR:** {len(filtered)}")
        
        if len(filtered) > 0:
            st.dataframe(filtered, use_container_width=True, hide_index=True)
        else:
            st.info("No guests found in this amount range.")

# ---------------- UPDATE GUEST ----------------
elif option == "Update Guest":
    st.subheader("✏️ Update Guest")
    
    selected = st.selectbox("Select a guest", df["Name"].tolist())
    idx = df[df["Name"] == selected].index[0]
    
    col1, col2 = st.columns(2)
    
    with col1:
        new_name = st.text_input("Update Name", df.at[idx, "Name"])
        new_city = st.text_input("Update City", df.at[idx, "City"])
    
    with col2:
        new_invited = st.checkbox("Invited ✅", df.at[idx, "Invited"])
        new_amount = st.number_input("Amount (PKR)", 
                                   value=int(df.at[idx, "Amount"]), 
                                   min_value=0, step=100)

    if st.button("Save Changes", type="primary"):
        if new_name.strip() == "":
            st.warning("⚠ Name cannot be empty!")
        else:
            # Create a copy of the dataframe and update it
            updated_df = df.copy()
            updated_df.at[idx, "Name"] = new_name.strip()
            updated_df.at[idx, "City"] = new_city.strip()
            updated_df.at[idx, "Invited"] = new_invited
            updated_df.at[idx, "Amount"] = new_amount
            save_data(updated_df)

# ---------------- REMOVE GUEST - FIXED VERSION ----------------
elif option == "Remove Guest":
    st.subheader("🗑 Remove Guest")
    
    if df.empty:
        st.info("📝 No guests available to remove.")
    else:
        selected = st.selectbox("Select guest to remove", df["Name"].tolist())
        
        if selected:
            guest_info = df[df["Name"] == selected].iloc[0]
            
            st.error("🚨 **PERMANENT DELETION**")
            st.warning(f"You are about to **permanently delete**: **{selected}**")
            
            # Display guest details
            col1, col2 = st.columns(2)
            with col1:
                st.write("**Guest Details:**")
                st.write(f"• City: {guest_info['City']}")
                st.write(f"• Amount: PKR {guest_info['Amount']:,}")
                st.write(f"• Invited: {'✅ Yes' if guest_info['Invited'] else '❌ No'}")
            
            with col2:
                st.write("**Deletion Impact:**")
                st.write("• ❌ Guest will be permanently removed")
                st.write("• ❌ Cannot be recovered")
                st.write("• 📊 Total guests will decrease by 1")
            
            # Double confirmation
            st.markdown("---")
            st.write("**Type the guest's exact name to confirm deletion:**")
            confirm_text = st.text_input("Confirmation", placeholder=f"Type: {selected}")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("🔥 PERMANENTLY DELETE", 
                           type="primary",
                           disabled=confirm_text != selected):
                    # Perform deletion - create new dataframe without the selected guest
                    updated_df = df[df["Name"] != selected].reset_index(drop=True)
                    save_data(updated_df)
                    st.success(f"✅ {selected} has been permanently deleted!")
            
            with col2:
                if st.button("❌ Cancel Deletion"):
                    st.info("Deletion cancelled - no changes made")

# ---------------- FOOTER ----------------
st.markdown("---")
st.markdown("<p style='text-align:center;color:gray;'>🖋 Prepared by <b>Susheel Kumar</b></p>", unsafe_allow_html=True)